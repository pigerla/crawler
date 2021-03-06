__author__ = 'spy'

import pymongo
from docx import Document
from crawler import settings

conn = pymongo.Connection(settings.MONGODB_SERVER ,settings.MONGODB_PORT)
db = conn.crawler
items = db.thjy.find()
document = Document()
for item in items:
    document.add_heading(item['title'] , level=1)
    document.add_paragraph(item['content'])

document.add_page_break()
document.save('siNianJiDuShuBiJiShang.doc')