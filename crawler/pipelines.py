# -*- coding: utf-8 -*-

# Define your item pipelines here
#
# Don't forget to add your pipeline to the ITEM_PIPELINES setting
# See: http://doc.scrapy.org/en/latest/topics/item-pipeline.html

from docx import Document


class CrawlerPipeline(object):
    def __init__(self):
        self.document = ''

    def process_item(self, item, spider):
        self.document.add_heading(item['title'] , level=1)
        self.document.add_paragraph(item['content'])

        return item
    def open_spider(self, spider):
        print spider.name, 'is opening'
        self.document = Document()

    def close_spider(self, spider):
        print spider.name, 'is closing'
        self.document.add_page_break()
        self.document.save('siNianJiDuShuBiJiShang.doc')
